import random
import string
import os
import subprocess
from copy import deepcopy
from openpyxl import load_workbook  # type: ignore
from docx import Document  # type: ignore
from docx.shared import Pt, Inches, RGBColor  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER  # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE  # type: ignore
from docx.enum.section import WD_ORIENT, WD_SECTION  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore

ROWS = 8
COLS = 12

BASE_FOLDER = r"C:\Users\riley\OneDrive\Desktop\wordsearch_booklet"

EXCEL_PATH = os.path.join(BASE_FOLDER, "Book2.xlsx")
PUZZLES_FOLDER = os.path.join(BASE_FOLDER, "Puzzles")
ANSWER_KEY_FOLDER = os.path.join(BASE_FOLDER, "Answer Key")
OUTPUT_FOLDER = BASE_FOLDER

TITLE_FONT_SIZE = 26
GRID_FONT_SIZE = 36
WORD_FONT_SIZE = 18

GRID_CELL_WIDTH = 58
GRID_ROW_HEIGHT = 42
WORD_BANK_ROW_HEIGHT = 24
WORD_BANK_COL_WIDTH = 155

LANDSCAPE_TOP_MARGIN = 0.25
LANDSCAPE_BOTTOM_MARGIN = 0.25
LANDSCAPE_LEFT_MARGIN = 0.25
LANDSCAPE_RIGHT_MARGIN = 0.25

PORTRAIT_TOP_MARGIN = 0.75
PORTRAIT_BOTTOM_MARGIN = 0.75
PORTRAIT_LEFT_MARGIN = 0.75
PORTRAIT_RIGHT_MARGIN = 0.75

ANSWER_TITLE_SIZE = 16
ANSWER_GRID_FONT_SIZE = 14
ANSWER_CELL_SIZE = 24
ANSWER_BLOCK_HEIGHT = 255


def verify_required_paths():
    required_folders = [
        PUZZLES_FOLDER,
        ANSWER_KEY_FOLDER,
    ]

    required_files = [
        EXCEL_PATH,
        os.path.join(PUZZLES_FOLDER, "Puzzles Cover.docx"),
        os.path.join(PUZZLES_FOLDER, "Puzzles Copyright.docx"),
        os.path.join(PUZZLES_FOLDER, "Puzzles Instructions.docx"),
        os.path.join(PUZZLES_FOLDER, "Puzzles Back Cover.docx"),
        os.path.join(ANSWER_KEY_FOLDER, "Answer Key Cover.docx"),
        os.path.join(ANSWER_KEY_FOLDER, "Answer Key Copyright.docx"),
        os.path.join(ANSWER_KEY_FOLDER, "Answer Key Back Cover.docx"),
    ]

    print("Checking required folders...")
    for folder in required_folders:
        print(f"  {folder}")
        if not os.path.isdir(folder):
            raise FileNotFoundError(f"Missing folder: {folder}")

    print("Checking required files...")
    for file_path in required_files:
        print(f"  {file_path}")
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"Missing file: {file_path}")


def load_puzzles_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    ws = wb.worksheets[0]

    puzzle_titles = []
    puzzles = []

    for excel_row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not row or row[0] is None:
            continue

        title = str(row[0]).strip()
        if title == "":
            continue

        words = []
        for cell in row[1:]:
            if cell is None:
                continue

            word = str(cell).strip().upper()
            if word == "":
                continue

            if len(word) > COLS:
                print(f"Skipping word '{word}' in row {excel_row_num}: too long for {COLS} columns")
                continue

            words.append((word, 1, 1))

        if len(words) > ROWS:
            print(f"Row {excel_row_num} has more than {ROWS} words. Extra words will be ignored.")
            words = words[:ROWS]

        if words:
            puzzle_titles.append(title)
            puzzles.append(words)

    return puzzle_titles, puzzles


def build_grid(words):
    grid = [["" for _ in range(COLS)] for _ in range(ROWS)]

    shuffled_rows = list(range(ROWS))
    random.shuffle(shuffled_rows)

    placed_words = []

    for index, (word, _, _) in enumerate(words):
        r = shuffled_rows[index]
        max_start = COLS - len(word)
        c = random.randint(0, max_start)

        for i, ch in enumerate(word):
            grid[r][c + i] = ch

        placed_words.append((word, r + 1, c + 1))

    for r in range(ROWS):
        for c in range(COLS):
            if grid[r][c] == "":
                grid[r][c] = random.choice(string.ascii_uppercase)

    return grid, placed_words


def get_answer_positions(words):
    positions = set()
    for word, r, c in words:
        r -= 1
        c -= 1
        for i in range(len(word)):
            if 0 <= r < ROWS and 0 <= c + i < COLS:
                positions.add((r, c + i))
    return positions


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)

    tblPr.append(borders)


def style_grid_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(GRID_FONT_SIZE)
            run.font.name = "Arial"


def style_word_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(WORD_FONT_SIZE)
            run.font.name = "Arial"


def add_title(doc, title_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run(f'"{title_text}"')
    run.bold = True
    run.font.size = Pt(TITLE_FONT_SIZE)
    run.font.name = "Arial Rounded MT Bold"


def add_spacer(doc, points_after):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points_after)


def clear_footer(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def unlink_footer_from_previous(section):
    section.footer.is_linked_to_previous = False


def ensure_pgNumType(sectPr):
    pg_num_type = sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sectPr.append(pg_num_type)
    return pg_num_type


def set_page_number_start(section, start_value):
    sectPr = section._sectPr
    pg_num_type = ensure_pgNumType(sectPr)
    pg_num_type.set(qn("w:start"), str(start_value))


def set_page_number_format(section, fmt):
    sectPr = section._sectPr
    pg_num_type = ensure_pgNumType(sectPr)
    pg_num_type.set(qn("w:fmt"), fmt)


def hide_footer_for_section(section):
    unlink_footer_from_previous(section)
    footer = section.footer
    if footer.paragraphs:
        clear_footer(footer.paragraphs[0])


def setup_footer_for_section(section, left_text, right_text):
    unlink_footer_from_previous(section)
    footer = section.footer
    paragraph = footer.paragraphs[0]

    clear_footer(paragraph)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.right_indent = Inches(0.3)

    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.75), WD_TAB_ALIGNMENT.CENTER, WD_TAB_LEADER.SPACES)
    tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    left_run = paragraph.add_run(left_text)
    left_run.font.name = "Arial"
    left_run.font.size = Pt(10)

    paragraph.add_run("\t")

    page_run = paragraph.add_run()
    page_run.font.name = "Arial"
    page_run.font.size = Pt(10)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    page_run._r.append(fld_char_begin)
    page_run._r.append(instr_text)
    page_run._r.append(fld_char_sep)
    page_run._r.append(fld_char_end)

    paragraph.add_run("\t")

    right_run = paragraph.add_run(right_text)
    right_run.font.name = "Arial"
    right_run.font.size = Pt(10)


def set_section_landscape(section, footer_left, footer_right):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(LANDSCAPE_LEFT_MARGIN)
    section.right_margin = Inches(LANDSCAPE_RIGHT_MARGIN)
    section.top_margin = Inches(LANDSCAPE_TOP_MARGIN)
    section.bottom_margin = Inches(LANDSCAPE_BOTTOM_MARGIN)
    setup_footer_for_section(section, footer_left, footer_right)


def set_section_portrait(section, footer_left="", footer_right=""):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(PORTRAIT_LEFT_MARGIN)
    section.right_margin = Inches(PORTRAIT_RIGHT_MARGIN)
    section.top_margin = Inches(PORTRAIT_TOP_MARGIN)
    section.bottom_margin = Inches(PORTRAIT_BOTTOM_MARGIN)
    if footer_left or footer_right:
        setup_footer_for_section(section, footer_left, footer_right)
    else:
        hide_footer_for_section(section)


def append_doc_contents(target_doc, source_path):
    src = Document(source_path)

    body = target_doc._element.body
    src_body = src._element.body

    for child in src_body:
        if child.tag.endswith("sectPr"):
            continue
        body.append(deepcopy(child))


def append_external_doc_as_own_page(target_doc, source_path):
    if len(target_doc._element.body) > 0:
        target_doc.add_section(WD_SECTION.NEW_PAGE)

    section = target_doc.sections[-1]
    set_section_portrait(section)
    append_doc_contents(target_doc, source_path)
    return section


def start_new_landscape_section(doc, footer_left, footer_right):
    if len(doc._element.body) > 0:
        doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    set_section_landscape(section, footer_left, footer_right)
    return section


def build_generated_puzzles_section(doc, puzzle_titles, puzzles):
    puzzle_section = start_new_landscape_section(
        doc,
        '"Dementia Friendly Wordsearch"',
        '"Puzzles by Riley"'
    )
    set_page_number_format(puzzle_section, "decimal")
    set_page_number_start(puzzle_section, 1)

    for i, puzzle in enumerate(puzzles, start=1):
        if i > 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
            next_section = doc.sections[-1]
            set_section_landscape(
                next_section,
                '"Dementia Friendly Wordsearch"',
                '"Puzzles by Riley"'
            )

        title = puzzle_titles[i - 1]
        grid, _placed_words = build_grid(puzzle)

        add_spacer(doc, 6)
        add_title(doc, title)
        add_spacer(doc, 8)

        table = doc.add_table(rows=ROWS, cols=COLS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(table)

        for row in table.rows:
            row.height = Pt(GRID_ROW_HEIGHT)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                cell.width = Pt(GRID_CELL_WIDTH)

        for r in range(ROWS):
            for c in range(COLS):
                cell = table.cell(r, c)
                cell.text = grid[r][c]
                style_grid_cell(cell)

        add_spacer(doc, 10)

        words_only = [word for word, _, _ in puzzle]
        word_bank = doc.add_table(rows=2, cols=4)
        word_bank.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(word_bank)

        for row in word_bank.rows:
            row.height = Pt(WORD_BANK_ROW_HEIGHT)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                cell.width = Pt(WORD_BANK_COL_WIDTH)

        index = 0
        for r in range(2):
            for c in range(4):
                cell = word_bank.cell(r, c)
                if index < len(words_only):
                    cell.text = words_only[index]
                    style_word_cell(cell)
                index += 1


def build_generated_answers_section(doc, puzzle_titles, puzzles):
    answer_keys = []
    for i, puzzle in enumerate(puzzles):
        title = puzzle_titles[i]
        grid, placed_words = build_grid(puzzle)
        answer_keys.append((title, grid, get_answer_positions(placed_words)))

    answer_section = start_new_landscape_section(
        doc,
        '"Dementia Friendly Wordsearch - Answer Key"',
        '"Puzzles by Riley"'
    )
    set_page_number_format(answer_section, "decimal")
    set_page_number_start(answer_section, 1)

    first_page = True
    for start in range(0, len(answer_keys), 4):
        if not first_page:
            doc.add_section(WD_SECTION.NEW_PAGE)
            next_section = doc.sections[-1]
            set_section_landscape(
                next_section,
                '"Dementia Friendly Wordsearch - Answer Key"',
                '"Puzzles by Riley"'
            )
        first_page = False

        block = answer_keys[start:start + 4]

        add_spacer(doc, 4)

        answer_page = doc.add_table(rows=2, cols=2)
        answer_page.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(answer_page)

        for row in answer_page.rows:
            row.height = Pt(ANSWER_BLOCK_HEIGHT)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        slot = 0
        for rr in range(2):
            for cc in range(2):
                outer_cell = answer_page.cell(rr, cc)

                if slot < len(block):
                    title, grid, positions = block[slot]

                    top_space = outer_cell.paragraphs[0]
                    top_space.paragraph_format.space_before = Pt(0)
                    top_space.paragraph_format.space_after = Pt(8)

                    title_p = outer_cell.add_paragraph()
                    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_p.paragraph_format.space_before = Pt(0)
                    title_p.paragraph_format.space_after = Pt(8)

                    title_run = title_p.add_run(f'Answer Key: "{title}"')
                    title_run.bold = True
                    title_run.font.size = Pt(ANSWER_TITLE_SIZE)
                    title_run.font.name = "Arial"

                    mini = outer_cell.add_table(rows=ROWS, cols=COLS)
                    mini.alignment = WD_TABLE_ALIGNMENT.CENTER
                    remove_table_borders(mini)

                    for row in mini.rows:
                        row.height = Pt(ANSWER_CELL_SIZE)
                        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                        for mini_cell in row.cells:
                            mini_cell.width = Pt(ANSWER_CELL_SIZE)

                    for r in range(ROWS):
                        for c in range(COLS):
                            mini_cell = mini.cell(r, c)
                            para = mini_cell.paragraphs[0]
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            para.paragraph_format.space_before = Pt(0)
                            para.paragraph_format.space_after = Pt(0)

                            run = para.add_run(grid[r][c])
                            run.font.size = Pt(ANSWER_GRID_FONT_SIZE)
                            run.font.name = "Arial"

                            if (r, c) in positions:
                                run.font.color.rgb = RGBColor(255, 0, 0)

                slot += 1


def save_document_safely(doc, file_path):
    base, ext = os.path.splitext(file_path)
    counter = 1

    while True:
        try:
            doc.save(file_path)
            return file_path
        except PermissionError:
            file_path = f"{base}_{counter}{ext}"
            counter += 1


def get_available_output_path(file_path):
    base, ext = os.path.splitext(file_path)
    candidate = file_path
    counter = 1

    while os.path.exists(candidate):
        candidate = f"{base}_{counter}{ext}"
        counter += 1

    return candidate


def open_folder(path):
    try:
        os.startfile(path)
    except Exception:
        try:
            subprocess.Popen(["explorer", path])
        except Exception:
            pass


def build_puzzles_book(puzzle_titles, puzzles):
    doc = Document()

    first_section = doc.sections[0]
    set_section_portrait(first_section)
    hide_footer_for_section(first_section)

    cover_section = append_external_doc_as_own_page(
        doc, os.path.join(PUZZLES_FOLDER, "Puzzles Cover.docx")
    )
    hide_footer_for_section(cover_section)

    copyright_section = append_external_doc_as_own_page(
        doc, os.path.join(PUZZLES_FOLDER, "Puzzles Copyright.docx")
    )
    set_section_portrait(copyright_section, '"Dementia Friendly Wordsearch"', '"Puzzles by Riley"')
    set_page_number_format(copyright_section, "lowerRoman")
    set_page_number_start(copyright_section, 1)

    instructions_section = append_external_doc_as_own_page(
        doc, os.path.join(PUZZLES_FOLDER, "Puzzles Instructions.docx")
    )
    set_section_portrait(instructions_section, '"Dementia Friendly Wordsearch"', '"Puzzles by Riley"')
    set_page_number_format(instructions_section, "lowerRoman")

    build_generated_puzzles_section(doc, puzzle_titles, puzzles)

    back_cover_section = append_external_doc_as_own_page(
        doc, os.path.join(PUZZLES_FOLDER, "Puzzles Back Cover.docx")
    )
    hide_footer_for_section(back_cover_section)

    return doc


def build_answers_book(puzzle_titles, puzzles):
    doc = Document()

    first_section = doc.sections[0]
    set_section_portrait(first_section)
    hide_footer_for_section(first_section)

    cover_section = append_external_doc_as_own_page(
        doc, os.path.join(ANSWER_KEY_FOLDER, "Answer Key Cover.docx")
    )
    hide_footer_for_section(cover_section)

    copyright_section = append_external_doc_as_own_page(
        doc, os.path.join(ANSWER_KEY_FOLDER, "Answer Key Copyright.docx")
    )
    set_section_portrait(copyright_section, '"Dementia Friendly Wordsearch - Answer Key"', '"Puzzles by Riley"')
    set_page_number_format(copyright_section, "lowerRoman")
    set_page_number_start(copyright_section, 1)

    build_generated_answers_section(doc, puzzle_titles, puzzles)

    back_cover_section = append_external_doc_as_own_page(
        doc, os.path.join(ANSWER_KEY_FOLDER, "Answer Key Back Cover.docx")
    )
    hide_footer_for_section(back_cover_section)

    return doc


def main():
    try:
        print("Starting wordsearch booklet generator...")
        print(f"Base folder: {BASE_FOLDER}")
        print(f"Excel file: {EXCEL_PATH}")
        print(f"Puzzles folder: {PUZZLES_FOLDER}")
        print(f"Answer key folder: {ANSWER_KEY_FOLDER}")
        print(f"Output folder: {OUTPUT_FOLDER}")
        print()

        verify_required_paths()
        os.makedirs(OUTPUT_FOLDER, exist_ok=True)

        puzzle_titles, puzzles = load_puzzles_from_excel(EXCEL_PATH)

        if not puzzles:
            print("No puzzles were found in the Excel file.")
            input("Press Enter to close...")
            return

        puzzles_doc = build_puzzles_book(puzzle_titles, puzzles)
        answers_doc = build_answers_book(puzzle_titles, puzzles)

        final_puzzles_path = get_available_output_path(
            os.path.join(OUTPUT_FOLDER, "wordsearch_puzzles_booklet.docx")
        )

        final_answers_path = get_available_output_path(
            os.path.join(OUTPUT_FOLDER, "wordsearch_answer_key_booklet.docx")
        )

        save_document_safely(puzzles_doc, final_puzzles_path)
        save_document_safely(answers_doc, final_answers_path)

        print()
        print("Saved puzzles booklet to:")
        print(final_puzzles_path)
        print()
        print("Saved answer key booklet to:")
        print(final_answers_path)
        print()
        print("Opening output folder...")
        open_folder(OUTPUT_FOLDER)

    except Exception as e:
        print()
        print("An error occurred:")
        print(str(e))
        print()
        print("Expected files:")
        print(EXCEL_PATH)
        print(os.path.join(PUZZLES_FOLDER, "Puzzles Cover.docx"))
        print(os.path.join(PUZZLES_FOLDER, "Puzzles Copyright.docx"))
        print(os.path.join(PUZZLES_FOLDER, "Puzzles Instructions.docx"))
        print(os.path.join(PUZZLES_FOLDER, "Puzzles Back Cover.docx"))
        print(os.path.join(ANSWER_KEY_FOLDER, "Answer Key Cover.docx"))
        print(os.path.join(ANSWER_KEY_FOLDER, "Answer Key Copyright.docx"))
        print(os.path.join(ANSWER_KEY_FOLDER, "Answer Key Back Cover.docx"))
        print()
        print("Also make sure required packages are installed:")
        print("pip install openpyxl python-docx")

    input("Press Enter to close...")


if __name__ == "__main__":
    main()